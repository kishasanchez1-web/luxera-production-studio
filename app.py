import streamlit as st
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from datetime import datetime
import tempfile

st.set_page_config(page_title="Luxera™ Production Studio", page_icon="✨", layout="wide")

st.markdown("""
<style>
.stApp { background-color:#F8F5EF; }
h1,h2,h3 { color:#0A0A0A; }
.block-container { max-width: 1200px; }
.lux-card { background:white; border:1px solid #C6A55A; padding:24px; border-radius:16px; }
</style>
""", unsafe_allow_html=True)

st.title("Luxera™ Production Studio")
st.caption("Generate luxury Luxera™ manuscript sections, DOCX files, and production-ready content.")

PILLARS = [
    "Executive Briefing™",
    "Family Office Insight™",
    "Hidden Risks™",
    "Hidden Opportunities™",
    "Guided Discovery™",
    "Decision Architecture™",
    "Wealth Architect Reflection™",
    "Executive Interpretation™",
    "Wealth Architect Advisory™",
    "Executive Decision™",
]

MODULES = [
    "Wealth Story™",
    "Future Self Wealth Blueprint™",
    "Opportunity Pipeline™",
    "Wealth Leak Detection™",
    "Family Office Annual Retreat™",
    "Financial Freedom Score™",
    "Financial Fragility Index™",
    "Debt Pressure Score™",
    "Wealth Velocity Score™",
    "Financial Confidence Index™",
    "Wealth Protection Framework™",
    "Freedom Lifestyle Design™",
    "Time Freedom Framework™",
    "Experience Portfolio™",
    "Net Worth Command Center™",
    "Wealth Architect Certification™",
]

DEFAULT_GUIDANCE = {
    "Wealth Story™": "Every financial decision is influenced by a story. Some stories create freedom. Some create limitation. This section uncovers the beliefs, experiences, victories, fears, and opportunities shaping the Wealth Architect’s current financial architecture.",
    "Future Self Wealth Blueprint™": "Most financial goals fail because they are disconnected from a meaningful vision. This section creates the detailed future that wealth is intended to support.",
    "Opportunity Pipeline™": "Opportunities often hide inside skills, relationships, assets, lived experience, technology, knowledge, and market demand. This section identifies and ranks opportunities by value, ease, speed, and freedom impact.",
    "Wealth Leak Detection™": "Most people focus on earning more. Few focus on recovering what already belongs to them. This section identifies money, resources, opportunities, and value currently escaping the financial architecture.",
    "Family Office Annual Retreat™": "This is the centerpiece of the Luxera™ experience: a private family-office-inspired annual strategy session to review, redesign, and recommit to freedom and legacy.",
}

def add_logo_or_title(doc, logo_path=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if logo_path and Path(logo_path).exists():
        try:
            p.add_run().add_picture(str(logo_path), width=Inches(2.1))
            return
        except Exception:
            pass
    r = p.add_run("Luxera™")
    r.font.size = Pt(32)
    r.bold = True
    r.font.color.rgb = RGBColor(198,165,90)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = "Georgia"
        r.font.color.rgb = RGBColor(10,10,10)
    return h

def add_p(doc, text="", bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Georgia"
    r.font.size = Pt(11)
    r.bold = bold
    return p

def add_lines(doc, count=10):
    for _ in range(count):
        add_p(doc, "_" * 76)

def build_section(doc, module_name, guidance, reflection_lines):
    doc.add_page_break()
    add_heading(doc, module_name, 1)
    add_p(doc, guidance)
    for pillar in PILLARS:
        add_heading(doc, pillar, 2)
        if pillar == "Executive Briefing™":
            add_p(doc, f"This {module_name} section is designed to move the Wealth Architect™ beyond surface answers and into strategic clarity. The objective is not to fill space. The objective is to identify the decisions, risks, opportunities, and priorities that shape financial freedom.")
        elif pillar == "Family Office Insight™":
            add_p(doc, "Private wealth strategy begins by asking better questions. The quality of the question determines the quality of the decision.")
        elif pillar == "Hidden Risks™":
            add_p(doc, "Unexamined assumptions, delayed decisions, unclear priorities, and inherited beliefs can quietly shape financial outcomes for years.")
        elif pillar == "Hidden Opportunities™":
            add_p(doc, "Opportunity is often already present. It may exist inside existing skills, relationships, assets, knowledge, technology, or underused resources.")
        elif pillar == "Guided Discovery™":
            add_p(doc, "Before answering, consider what is already working, what is creating friction, what would improve quality of life, and what future self would insist you prioritize.")
        elif pillar == "Decision Architecture™":
            add_p(doc, "Evaluate each possible action through five filters: Does it increase freedom? Reduce risk? Create opportunity? Improve quality of life? Support legacy?")
        elif pillar == "Wealth Architect Reflection™":
            add_p(doc, "Use the space below for a complete answer. Do not rush this section.")
            add_lines(doc, reflection_lines)
        elif pillar == "Executive Interpretation™":
            add_p(doc, "Your answers reveal patterns. Look for repeated themes, emotional intensity, avoidance, excitement, and friction. Those signals often identify the highest-value next action.")
        elif pillar == "Wealth Architect Advisory™":
            add_p(doc, "Prioritize the action that creates the greatest improvement across multiple categories. The best action usually improves freedom, confidence, protection, and opportunity simultaneously.")
        elif pillar == "Executive Decision™":
            add_p(doc, "The single highest-value action I will take is:")
            add_lines(doc, 4)

def create_doc(selected, logo_file, reflection_lines):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)

    logo_path = None
    if logo_file:
        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=Path(logo_file.name).suffix)
        tmp.write(logo_file.read())
        tmp.close()
        logo_path = tmp.name

    add_logo_or_title(doc, logo_path)
    add_heading(doc, "Financial Freedom Operating System™", 0)
    p = doc.add_paragraph("A Luxury Personal Wealth Architecture System™")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "Confidential Production Edition", 1)
    add_p(doc, "This document was generated by Luxera™ Production Studio using the luxury advisory framework: Executive Briefing™, Family Office Insight™, Hidden Risks™, Hidden Opportunities™, Guided Discovery™, Decision Architecture™, Wealth Architect Reflection™, Executive Interpretation™, Wealth Architect Advisory™, and Executive Decision™.")

    for m in selected:
        build_section(doc, m, DEFAULT_GUIDANCE.get(m, "This section develops strategic clarity inside the Luxera™ Financial Freedom Operating System™."), reflection_lines)

    out = Path(tempfile.gettempdir()) / f"Luxera_Production_Studio_Output_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(out)
    return out

st.sidebar.header("Build Settings")
selected = st.sidebar.multiselect("Select modules", MODULES, default=["Wealth Story™", "Future Self Wealth Blueprint™", "Opportunity Pipeline™", "Wealth Leak Detection™", "Family Office Annual Retreat™"])
reflection_lines = st.sidebar.slider("Writing lines per reflection", 8, 28, 16)
logo_file = st.sidebar.file_uploader("Upload Luxera™ logo PNG/JPG", type=["png", "jpg", "jpeg"])

st.markdown('<div class="lux-card">', unsafe_allow_html=True)
st.subheader("Production Builder")
st.write("Select modules, upload the Luxera™ logo if desired, then generate a DOCX.")
if st.button("Generate Luxera™ DOCX"):
    output = create_doc(selected, logo_file, reflection_lines)
    st.success("Luxera™ DOCX generated.")
    st.download_button("Download DOCX", output.read_bytes(), file_name=output.name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
st.markdown('</div>', unsafe_allow_html=True)
